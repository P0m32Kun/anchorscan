"""Build a visible-slot docxtpl prototype from the immutable source copy."""

from __future__ import annotations

import copy
import json
import shutil
import struct
import unicodedata
import zlib
import zipfile
from copy import deepcopy
from pathlib import Path

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from jinja2 import Environment
from lxml import etree

from check_structure import verify


ROOT = Path(__file__).parent
FIXTURES = ROOT / "fixtures"
OUT = ROOT / "out"
SOURCE = FIXTURES / "source-template.docx"
TEMPLATE = OUT / "project-report-placeholder-template.docx"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W, "wp": WP, "a": A, "r": R}


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def write_png(path: Path, width: int, height: int, rgb: tuple[int, int, int]) -> None:
    raw = b"".join(b"\0" + bytes(rgb) * width for _ in range(height))
    chunk = lambda kind, data: (
        struct.pack(">I", len(data))
        + kind
        + data
        + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
    )
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw, 9))
        + chunk(b"IEND", b"")
    )


def node_text(node: etree._Element) -> str:
    return "".join(node.xpath(".//w:t/text()", namespaces=NS)).strip()


def cover_line(value: object, width: float = 52) -> str:
    text = str(value)
    text_width = sum(
        3.5 if unicodedata.east_asian_width(char) in "WFA" else 3 for char in text
    )
    padding = max(0, round(width - text_width))
    left = padding // 2
    return "\u00a0" * left + text + "\u00a0" * (padding - left)


def set_row_text(row: etree._Element, values: list[str]) -> None:
    for cell, value in zip(row.findall("w:tc", NS), values, strict=True):
        texts = cell.xpath(".//w:t", namespaces=NS)
        if not texts:
            paragraph_node = cell.find("w:p", NS)
            run = etree.SubElement(paragraph_node, w("r"))
            texts = [etree.SubElement(run, w("t"))]
        texts[0].text = value
        for text in texts[1:]:
            text.text = ""


def merge_row(row: etree._Element) -> None:
    cells = row.findall("w:tc", NS)
    properties = cells[0].find("w:tcPr", NS)
    if properties is None:
        properties = etree.SubElement(cells[0], w("tcPr"))
    span = properties.find("w:gridSpan", NS)
    if span is None:
        span = etree.SubElement(properties, w("gridSpan"))
    span.set(w("val"), str(len(cells)))
    for cell in cells[1:]:
        row.remove(cell)


def paragraph(text: str) -> etree._Element:
    node = etree.Element(w("p"))
    run = etree.SubElement(node, w("r"))
    etree.SubElement(run, w("t")).text = text
    return node


def with_text(sample: etree._Element, text: str) -> etree._Element:
    node = deepcopy(sample)
    runs = node.findall("w:r", NS)
    replacement = deepcopy(runs[0]) if runs else etree.Element(w("r"))
    for child in list(replacement):
        if child.tag != w("rPr"):
            replacement.remove(child)
    etree.SubElement(replacement, w("t")).text = text
    for child in list(node):
        if child.tag != w("pPr"):
            node.remove(child)
    node.append(replacement)
    return node


def without_numbering(paragraph_node: etree._Element) -> etree._Element:
    properties = paragraph_node.find("w:pPr", NS)
    existing = properties.find("w:numPr", NS)
    if existing is not None:
        properties.remove(existing)
    numbering = etree.Element(w("numPr"))
    etree.SubElement(numbering, w("numId")).set(w("val"), "0")
    style = properties.find("w:pStyle", NS)
    properties.insert(list(properties).index(style) + 1, numbering)
    return paragraph_node


def replace_paragraph(body: etree._Element, old_text: str, new_text: str) -> None:
    matches = [p for p in body.xpath("./w:p", namespaces=NS) if node_text(p) == old_text]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph {old_text!r}, got {len(matches)}")
    replacement = with_text(matches[0], new_text)
    body.replace(matches[0], replacement)


def replace_labeled_slot(
    body: etree._Element,
    old_text: str,
    label: str,
    slot: str,
    value_run_sample: etree._Element,
) -> None:
    matches = [p for p in body.xpath("./w:p", namespaces=NS) if node_text(p) == old_text]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph {old_text!r}, got {len(matches)}")
    source = matches[0]
    replacement = deepcopy(source)
    for child in list(replacement):
        if child.tag != w("pPr"):
            replacement.remove(child)

    properties = replacement.find("w:pPr", NS)
    existing_tabs = properties.find("w:tabs", NS)
    if existing_tabs is not None:
        properties.remove(existing_tabs)

    source_runs = source.findall("w:r", NS)
    label_run = deepcopy(source_runs[0])
    for child in list(label_run):
        if child.tag != w("rPr"):
            label_run.remove(child)
    etree.SubElement(label_run, w("t")).text = label
    replacement.append(label_run)

    value_run = deepcopy(value_run_sample)
    for child in list(value_run):
        if child.tag != w("rPr"):
            value_run.remove(child)
    underline = value_run.find("w:rPr/w:u", NS)
    if underline is None:
        underline = etree.SubElement(value_run.find("w:rPr", NS), w("u"))
    underline.set(w("val"), "single")
    etree.SubElement(value_run, w("t")).text = slot
    replacement.append(value_run)
    body.replace(source, replacement)


def prepare_summary_table(body: etree._Element) -> None:
    tables = body.findall("w:tbl", NS)
    if len(tables) != 1:
        raise RuntimeError(f"expected one summary table, got {len(tables)}")
    table = tables[0]
    rows = table.findall("w:tr", NS)
    if len(rows) != 2:
        raise RuntimeError(f"expected header plus one data row, got {len(rows)}")
    data = rows[1]
    additions = []
    for values, merged in (
        (["{%tr for r in summary_rows %}", "", "", ""], True),
        (["{{ r.no }}", "{{ r.title }}", "{{ r.assets }}", "{{ r.level }}"], False),
        (["{%tr endfor %}", "", "", ""], True),
        (["{%tr if summary_empty %}", "", "", ""], True),
        (["本次纳入报告范围内无已确认漏洞", "", "", ""], True),
        (["{%tr endif %}", "", "", ""], True),
    ):
        row = deepcopy(data)
        set_row_text(row, values)
        if merged:
            merge_row(row)
        additions.append(row)
    index = list(table).index(data)
    table.remove(data)
    for offset, row in enumerate(additions):
        table.insert(index + offset, row)


def prepare_network_zone_block(body: etree._Element, untouched: dict[str, bytes]) -> None:
    children = list(body)
    texts = [node_text(node) for node in children]
    zone_start = texts.index("I区")
    conclusion = texts.index("渗透测试结论")
    old_zone_nodes = children[zone_start:conclusion]
    if any(node.xpath(".//w:sectPr", namespaces=NS) for node in old_zone_nodes):
        raise RuntimeError("refusing to replace a network-zone range containing sectPr")

    zone_heading = children[zone_start]
    normal = children[zone_start + 1]
    heading3 = next(node for node in old_zone_nodes if node_text(node) == "（高危）")
    heading4 = next(node for node in old_zone_nodes if node_text(node) == "漏洞描述")
    source_image_ids = {
        blip.get(f"{{{R}}}embed")
        for node in old_zone_nodes
        for blip in node.xpath(".//a:blip", namespaces=NS)
    }

    for node in old_zone_nodes:
        body.remove(node)

    relationships = etree.fromstring(untouched["word/_rels/document.xml.rels"])
    removed_targets = []
    for relationship in list(relationships):
        if relationship.get("Id") in source_image_ids:
            removed_targets.append(relationship.get("Target"))
            relationships.remove(relationship)
    if len(removed_targets) != len(source_image_ids):
        raise RuntimeError("could not remove every source body image relationship")
    untouched["word/_rels/document.xml.rels"] = etree.tostring(
        relationships, encoding="UTF-8", xml_declaration=True
    )
    for target in removed_targets:
        untouched.pop(f"word/{target}")

    p = lambda text: with_text(normal, text)
    h3 = lambda text: with_text(heading3, text)
    h4 = lambda text: with_text(heading4, text)
    zone_nodes = [
        paragraph("{%p for network_zone in network_zones %}"),
        with_text(zone_heading, "{{ network_zone.name }}"),
        paragraph("{%p for session in network_zone.sessions %}"),
        p("接入记录：{{ session.label }}"),
        p("接入点：{{ session.access_point }}"),
        p("测试机 IP：{{ session.tester_ip }}"),
        p("测试范围：{{ session.targets_text }}"),
        paragraph("{%p if session.exclusions_text %}"),
        p("排除范围：{{ session.exclusions_text }}"),
        paragraph("{%p endif %}"),
        paragraph("{%p if session.notes %}"),
        p("备注：{{ session.notes }}"),
        paragraph("{%p endif %}"),
        paragraph("{%p endfor %}"),
        paragraph("{%p for verification in network_zone.confirmed %}"),
        h3("{{ verification.heading }}"),
        h4("漏洞描述"),
        p("{{ verification.description }}"),
        h4("漏洞详情"),
        paragraph("{%p for evidence in verification.evidence %}"),
        p("{{ evidence.image }}"),
        paragraph("{%p endfor %}"),
        h4("关联资产"),
        p("{{ verification.assets_text }}"),
        h4("修改建议"),
        p("{{ verification.remediation }}"),
        paragraph("{%p endfor %}"),
        paragraph("{%p if network_zone.not_observed %}"),
        without_numbering(h3("{{ network_zone.name }}其它漏洞验证不存在的截图：")),
        paragraph("{%p for verification in network_zone.not_observed %}"),
        p("{{ verification.title }}不存在证明，端口（{{ verification.ports_text }}）"),
        paragraph("{%p for evidence in verification.evidence %}"),
        p("{{ evidence.image }}"),
        paragraph("{%p endfor %}"),
        paragraph("{%p endfor %}"),
        paragraph("{%p endif %}"),
        paragraph("{%p endfor %}"),
    ]
    for offset, node in enumerate(zone_nodes):
        body.insert(zone_start + offset, node)


def prepare_conclusion(body: etree._Element) -> None:
    children = list(body)
    heading_index = [node_text(node) for node in children].index("渗透测试结论")
    old_paragraphs = [node for node in children[heading_index + 1 :] if node.tag == w("p")]
    if len(old_paragraphs) < 2:
        raise RuntimeError(f"expected at least two conclusion paragraphs, got {len(old_paragraphs)}")
    sample = old_paragraphs[0]
    for node in old_paragraphs:
        body.remove(node)
    conclusion_nodes = [
        with_text(
            sample,
            "本次测试共测试{{ report.client_name }}{{ conclusion.network_zone_names_text }}所有设备，"
            "共发现漏洞{{ conclusion.total }}个，其中严重漏洞{{ conclusion.critical }}个、高危漏洞{{ conclusion.high }}个、中危漏洞"
            "{{ conclusion.medium }}个、低危漏洞{{ conclusion.low }}个。其中问题主要集中在"
            "{{ conclusion.focus_text }}这几个方面。",
        ),
        with_text(sample, "需要及时整改，加强安全管理规范，并进行复测，预防安全事故发生。"),
    ]
    for node in conclusion_nodes:
        body.insert(len(body) - 1, node)


def rebuild_toc_field(document: etree._Element) -> None:
    """Replace the TOC field's cached result so it no longer references bookmarks.

    The source template hard-codes three zone entries (I/II/III) whose PAGEREF
    targets are deleted once zones become a dynamic loop. We keep the TOC field
    skeleton (begin/instr/separate ... end) with an empty result plus a hint, so
    Word/WPS rebuilds the entries from the live Heading paragraphs on open
    (updateFields is already set).
    """
    body = document.find("w:body", NS)
    children = list(body)
    toc_begin_idx = None
    for i, node in enumerate(children):
        instr = "".join(t.text or "" for t in node.iter(w("instrText")))
        if "TOC" in instr and node.findall(".//w:fldChar[@w:fldCharType='begin']", NS):
            toc_begin_idx = i
            break
    if toc_begin_idx is None:
        return

    depth = 0
    toc_end_idx = None
    for j in range(toc_begin_idx, len(children)):
        node = children[j]
        if node.tag != w("p"):
            continue
        depth += len(node.findall(".//w:fldChar[@w:fldCharType='begin']", NS))
        depth -= len(node.findall(".//w:fldChar[@w:fldCharType='end']", NS))
        if j > toc_begin_idx and depth == 0:
            toc_end_idx = j
            break
    if toc_end_idx is None:
        return

    begin_para = children[toc_begin_idx]
    properties = begin_para.find("w:pPr", NS)
    for child in list(begin_para):
        if child is not properties:
            begin_para.remove(child)

    def add_run(*xml_children: etree._Element) -> None:
        run = etree.Element(w("r"))
        for child in xml_children:
            run.append(child)
        begin_para.append(run)

    begin = etree.Element(w("fldChar"))
    begin.set(w("fldCharType"), "begin")
    add_run(begin)
    instr = etree.Element(w("instrText"))
    instr.text = ' TOC \\o "1-2" \\h \\u '
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    add_run(instr)
    separate = etree.Element(w("fldChar"))
    separate.set(w("fldCharType"), "separate")
    add_run(separate)
    hint = etree.Element(w("t"))
    hint.text = "右键此处选择「更新域」以生成目录"
    add_run(hint)

    for node in children[toc_begin_idx + 1 : toc_end_idx]:
        body.remove(node)


def prepare_template() -> None:
    """Rewrite only the copied experimental template; never touch the user source."""
    shutil.copyfile(SOURCE, TEMPLATE)
    with zipfile.ZipFile(TEMPLATE) as archive:
        document = etree.fromstring(archive.read("word/document.xml"))
        settings = etree.fromstring(archive.read("word/settings.xml"))
        untouched = {
            name: archive.read(name)
            for name in archive.namelist()
            if name not in {"word/document.xml", "word/settings.xml"}
        }

    body = document.find("w:body", NS)
    prepare_summary_table(body)
    for old, new in (
        ("XX电力有限公司安全渗透测试分析报告", "{{ report.title }}"),
        ("二零二二年X月", "{{ report.project_created_month }}"),
        (
            "根据XX电力有限公司的要求，对XX电力有限公司多个信息系统进行系统安全渗透测试。"
            "本次渗透测试的时间安排为2022年X月X日，主要以人工检测的方法进行。",
            "根据{{ report.client_name }}的要求，对{{ report.test_subject }}开展安全渗透测试，"
            "测试时间为{{ report.test_period }}。本报告汇总纳入报告范围内的测试与人工验证结果。",
        ),
    ):
        replace_paragraph(body, old, new)
    month_paragraph = next(
        p
        for p in body.xpath("./w:p", namespaces=NS)
        if node_text(p) == "{{ report.project_created_month }}"
    )
    month_paragraph.find("w:pPr/w:rPr/w:rFonts", NS).set(w("hint"), "eastAsia")
    time_paragraph = next(
        p for p in body.xpath("./w:p", namespaces=NS) if node_text(p) == "测试时间：           2022年X月X日"
    )
    time_value_run = next(
        run for run in time_paragraph.findall("w:r", NS) if run.find("w:rPr/w:u", NS) is not None
    )
    for old, label, slot in (
        (
            "测试对象：            XX电力有限公司",
            "测试对象：",
            "{{ report.test_subject|cover_line }}",
        ),
        (
            "测试时间：           2022年X月X日",
            "测试时间：",
            "{{ report.project_created_date|cover_line }}",
        ),
        (
            "测试人员：                 XX",
            "测试人员：",
            "{{ report.testers_text|cover_line }}",
        ),
    ):
        replace_labeled_slot(body, old, label, slot, time_value_run)
    prepare_network_zone_block(body, untouched)
    prepare_conclusion(body)
    rebuild_toc_field(document)

    if not settings.xpath("w:updateFields", namespaces=NS):
        update = etree.SubElement(settings, w("updateFields"))
        update.set(w("val"), "true")

    with zipfile.ZipFile(TEMPLATE, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in untouched.items():
            archive.writestr(name, content)
        archive.writestr(
            "word/document.xml", etree.tostring(document, encoding="UTF-8", xml_declaration=True)
        )
        archive.writestr(
            "word/settings.xml", etree.tostring(settings, encoding="UTF-8", xml_declaration=True)
        )


def evidence_paths() -> dict[str, Path]:
    evidence = OUT / "evidence"
    evidence.mkdir(parents=True, exist_ok=True)
    paths = {
        "evidence-01-wide": evidence / "evidence-01-wide.png",
        "evidence-02-tall": evidence / "evidence-02-tall.png",
        "evidence-03-wide": evidence / "evidence-03-wide.png",
    }
    write_png(paths["evidence-01-wide"], 640, 320, (228, 74, 48))
    write_png(paths["evidence-02-tall"], 320, 640, (48, 125, 228))
    write_png(paths["evidence-03-wide"], 640, 320, (61, 159, 100))
    return paths


def image_box(path: Path, max_width_mm: float = 150, max_height_mm: float = 180) -> tuple[float, float]:
    data = path.read_bytes()
    if not data.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError(f"prototype fixture only supports PNG: {path}")
    width, height = struct.unpack(">II", data[16:24])
    scale = min(max_width_mm / width, max_height_mm / height)
    return width * scale, height * scale


def render(context: dict, destination: Path, images: dict[str, Path]) -> None:
    template = DocxTemplate(TEMPLATE)
    context = copy.deepcopy(context)
    context["summary_empty"] = not context["summary_rows"]
    context["summary_rows"] = [
        {
            "no": row["number"],
            "title": row["title"],
            "assets": row["assets_text"],
            "level": row["severity_label"],
        }
        for row in context["summary_rows"]
    ]
    for zone in context["network_zones"]:
        for key in ("confirmed", "not_observed"):
            for verification in zone[key]:
                for evidence in verification["evidence"]:
                    path = images[evidence["image"]]
                    width, height = image_box(path)
                    evidence["image"] = InlineImage(
                        template, str(path), width=Mm(width), height=Mm(height)
                    )
    environment = Environment(autoescape=True)
    environment.filters["cover_line"] = cover_line
    template.render(context, jinja_env=environment, autoescape=True)
    template.save(destination)


def main() -> None:
    OUT.mkdir(exist_ok=True)
    prepare_template()
    context = json.loads((FIXTURES / "project_report.json").read_text())
    images = evidence_paths()
    for count in (0, 1, len(context["summary_rows"])):
        case = copy.deepcopy(context)
        case["summary_rows"] = case["summary_rows"][:count]
        render(case, OUT / f"table-{count}.docx", images)
    shutil.copyfile(
        OUT / f"table-{len(context['summary_rows'])}.docx", OUT / "project-report.docx"
    )
    verify(SOURCE, TEMPLATE, OUT)
    print(f"template: {TEMPLATE}")
    print(f"fixture render: {OUT / 'project-report.docx'}")


if __name__ == "__main__":
    main()
